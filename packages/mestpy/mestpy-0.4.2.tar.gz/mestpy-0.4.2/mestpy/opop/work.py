def interactive_redact(directory):
    import os
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return
    lnames=[]
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        lnames.append(f_name)
    joined_string = "\n".join(lnames)
    from IPython.display import display, clear_output
    import re
    import numpy as np
    import copy
    from ipywidgets import widgets
    """ redact Textarea Widget interactive"""   
    allFilesName=widgets.Textarea()

    old_value = joined_string # только один раз будет joined
    allFilesName.value=old_value # изменяет глобальную
    allFilesName.disabled=True
    allFilesName.layout=layout = { 'width' : '50%' ,'height' : '250Px' }  
    
    ############################################################## def1    
    def multi_interactive_checkbox_widget(wTextarea):
        """ redact Textarea Widget value with checkboxes """
        names = ["Числа", "Скобки", "Точки", "Пробелы"]
        import re

        checkboxes = [widgets.Checkbox(value=False, description=label) for label in names]
        arg_dict = {names[i]: checkbox for i, checkbox in enumerate(checkboxes)}
        output = widgets.VBox(children=checkboxes)

        selected_data = []
        selected_bool=[]
        patterns=['\d',r"[(){}[\]]+",'\.',' ']

        def select_data(**kwargs):
            text = old_value
            selected_data.clear()
            selected_bool.clear()
            for key in kwargs:
                selected_bool.append(kwargs[key])
                if kwargs[key] is True:
                    selected_data.append(key)
            text1=text
            for i in range(0, len(checkboxes)):
                if selected_bool[i]:
                    pattern=patterns[i]
                    repl=''
                    text1=re.sub(pattern, repl, text1)
                    text1="\n".join([txt.strip() for txt in text1.split('\n')])
            wTextarea.value=text1
        out = widgets.interactive_output(select_data, arg_dict)

        def on_VBox_change(value):
            pass
        output.children[0].observe(on_VBox_change, 'value')
        
        return output, out

    ############################################################## def2
    def two_textbox_widget(wTextarea):
        import re
        """ Widget with a search field and a rename field  """

        def on_text_change(change):
            search_input = change['new']
            if search_input == '':
                wTextarea.value=old_value
            else:
                text=wTextarea.value
                repl=''
                pattern=search_input
                text1=re.sub(pattern, repl, text)
                text1="\n".join([txt.strip() for txt in text1.split('\n')])
                wTextarea.value=text1
        outputWText = widgets.Text()
        outputWText.observe(on_text_change, names='value')
        return outputWText

    ############################################################## def3
    def three_widgets(wTextarea): 
        """ Widget with insert string and a rename field  
            0 - in start position
                -1 - in end position
            else - after number posotion
        """
        import copy
        output = widgets.Text()  
        outDrop = widgets.Dropdown(
        placeholder='После позиции',
        options=['В начале', 'В конце', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12'],
        value='В начале',
        description='Место вставки:',
        ensure_option=True,
        disabled=False
        )

        def insertStrForFilesName(files_names,string,afterNumber):
            new_value=[]

            import re
            lnamesTemp=old_value.split('\n')
            for ids,f_name in enumerate(lnamesTemp):
                text=f_name
                pattern=""
                m=""
                repl=str(string)
                if afterNumber==0:
                    pattern=r"^" #в начале строки
                    m=re.sub(pattern, repl, text)
                elif afterNumber==-1 or afterNumber>len(f_name):
                    pattern=r"$" #в конце строки
                    m=re.sub(pattern, repl, text)
                else:
                    if int(afterNumber) < len(f_name):
                        m=f_name[:int(afterNumber)]+repl+f_name[int(afterNumber):]
                    else:
                        pattern=r"$" #в конце строки
                        m=re.sub(pattern, repl, text)
                new_name=m.strip()
                new_value.append(new_name)
            return new_value

        def on_text_change(change):
            search_input = change['new']
            if search_input == '':
                # Reset search field
                wTextarea.value=old_value
            else:
                string=output.value
                stri=outDrop.value

                if stri=='В начале':
                    afterNumber=0
                elif stri=='В конце':
                    afterNumber=-1
                else:
                    afterNumber=int(stri)

                text1=insertStrForFilesName(old_value,string,afterNumber)
                text1="\n".join([txt.strip() for txt in text1])
                wTextarea.value=text1
        output.observe(on_text_change, names='value')
        return output, outDrop
    
    ############################################################## def4
    def check_of_empty_list(func):
        def wrapper(lst):
            if lst:
                return func(lst)
            else:
                raise Exception("Empty input data")
        return wrapper
    ############################################################## def5
    def check_of_identical_name(func):
        def wrapper(lst):
            setlst = set(lst)
            if len(lst) == len(setlst):
                return func(lst)
            else:
                raise Exception("В списке есть одинаковые")
        return wrapper
    
    
    ####################################################### def6
    #Главная кнопка для сохранения файлов после изменения
    saveButton = widgets.Button(
        description='Файлы не проверны',
        disabled=True,
        layout = { 'width' : 'max-content' }
    )

    @saveButton.on_click
    def on_saveButton_clicked(b):
        import os
        files = os.listdir(directory)
        lst=allFilesName.value.split('\n')
        if len(lst)!=len(files):
            print("Количество файлов и имен не равны")
            return
        for ids,f in enumerate(files):
            f_name, f_ext=os.path.splitext(f)
            f_name_new = str(lst[ids])
            new_name='{}{}'.format(f_name_new, f_ext)
            os.rename(directory+'/'+f, directory+'/'+new_name)
        print('Файлы сохранены')
        saveButton.disabled=True
        toExcelButton.disabled=True
        saveButton.description='Файлы не проверны'

    ############################################################## def7
    #кнопка для проверки повторов имен файлоа
    checkButton = widgets.Button(
        description="Проверка корректности имен файлов",
        layout = { 'width' : 'max-content' }
    )

    @checkButton.on_click
    def on_checkButton_clicked(b):
        lst=allFilesName.value.split('\n')
        saveButton.disabled=True
        saveButton.description='Файлы не проверны'
        toExcelButton.disabled=True

        @check_of_empty_list
        @check_of_identical_name
        def checkList(lst):
            saveButton.disabled=False
            saveButton.description='Файлы проверены, можно сохранить'
            toExcelButton.disabled=False
        checkList(lst)

    ############################################################## def8
    #Кнопка для сохранения изменений аккордиона. Файлы не сохраняет
    saveAccordionChange = widgets.Button(
        description="После изменений элементов списка нажмите",
        layout = { 'width' : 'max-content' }
    )
    @saveAccordionChange.on_click
    def on_saveButton_clicked(b):
        old_value = allFilesName.value

    ############################################################## def9
    #Кнопка для редактирования вручную
    redactButton = widgets.Button(
        description="редактировать вручную",
        layout = { 'width' : 'max-content' }
    )

    @redactButton.on_click
    def redactButton_clicked(b):
         allFilesName.disabled = False  
            
    ############################################################## def10
    myDrop = widgets.Dropdown(
        options=['Отменить ВСЕ','1,2,3,4,...', '1 ,2 ,3 ,4 ,...', '01,02,03,04,...', '01 ,02 ,03 ,04 ,...', \
                 '001,002,003,004,...', '001 ,002 ,003 ,004 ,...'],
        value='Отменить ВСЕ',
        description='Индексы',
        disabled=False,
    )          
    def dropdown_eventhandler(change):
        lnamesTemp=allFilesName.value.strip().split('\n')
        if(myDrop.index == 1):
            val=[str(ids+1)+x for ids, x in enumerate(lnamesTemp) ]
            allFilesName.value="\n".join(val)
        elif(myDrop.index == 2):
            val=[str(ids+1)+' '+x for ids, x in enumerate(lnamesTemp) ]
            allFilesName.value="\n".join(val)
        elif(myDrop.index == 3):
            val=[str(ids+1).zfill(2)+x for ids, x in enumerate(lnamesTemp) ]
            allFilesName.value="\n".join(val)
        elif(myDrop.index == 4):
            val=[str(ids+1).zfill(2)+' '+x for ids, x in enumerate(lnamesTemp) ]
            allFilesName.value="\n".join(val)
        elif(myDrop.index == 5):
            val=[str(ids+1).zfill(3)+x for ids, x in enumerate(lnamesTemp) ]
            allFilesName.value="\n".join(val)
        elif(myDrop.index == 6):
            val=[str(ids+1).zfill(3)+' '+x for ids, x in enumerate(lnamesTemp) ]
            allFilesName.value="\n".join(val)
        else:
            #allFilesName.value=joined_string
            allFilesName.value=old_value
    myDrop.observe(dropdown_eventhandler, names='value')    
    
######################################################################
    output, out = multi_interactive_checkbox_widget(allFilesName)
    outText=two_textbox_widget(allFilesName)

    output1, drop=three_widgets(allFilesName)
    myThree_widgetsVBox=widgets.VBox(
        [drop, output1]
    )
######################################################################    
    
    ############################################################## def11  
    def subWithRE(pattern,repl):
        if not repl:
            new_txt=""
        if not pattern:
            print('Пусто')
            new_txt=old_value
        else:
            text = allFilesName.value
            lnames=text.split('\n')
            lnames_new=[re.sub(pattern, repl, x).strip() for x in lnames] 
            new_txt="\n".join(lnames_new)
        return new_txt
###########################################################################
    myText1=widgets.Text()
    myText2=widgets.Text()
########################################################################

    ############################################################## def12
    myFixButton=widgets.Button(
        description="Заменить"
    )
    def myFixButton_onclick(b):
        pattern=myText1.value
        repl=myText2.value
        allFilesName.value=subWithRE(pattern,repl)
    myFixButton.on_click(myFixButton_onclick)

    ############################################################## def13
    myUndoButton=widgets.Button(
        description="Отменить"
    )
    def myUndoButton_onclick(b):
        allFilesName.value=old_value
    myUndoButton.on_click(myUndoButton_onclick)

    
    myVBox=widgets.VBox([myText1, myText2,myFixButton,myUndoButton])
    accordion = widgets.Accordion(children=[myVBox])
    accordion.set_title(0, 'Заменить')
    myHBox=widgets.HBox([accordion])

    
    ############################################################## def13
    myAcc =  widgets.Accordion(
        layout = { 'width' : '350px' ,'height' : '240px' },
        children=[output, outText, myDrop, myThree_widgetsVBox,myHBox ],
        selected_index=None
    )
    myAcc.set_title(0, "Удалить лишние символы")
    myAcc.set_title(1, "Удалить слова")
    myAcc.set_title(2, "Добавить индексы")
    myAcc.set_title(3, "Добавить слово")
    myAcc.set_title(4, "Заменить слово")

    countPrint=1
    
    @myAcc.observe
    def myAcc_clicked(b):
        nonlocal countPrint
        countPrint+=1
        if countPrint%3==0:
            if myAcc.selected_index==0:
                myAcc.children[0].children[0].value = False
        sliderButton.disabled=False
        if len(cont1.children)==1:
            pass
        else:
            remove = cont1.children[-1]
            cont1.children = cont1.children[:-1]
            remove.close()
######################################################################## 
    myHBox = widgets.HBox([allFilesName,widgets.VBox([saveAccordionChange, myAcc])])
    display(myHBox)
###########################################################################

    ############################################################## def14
    #Кнопка для слайдера
    sliderButton = widgets.Button(
        description="укоротить имена файлов",
        layout = { 'width' : 'max-content' }
    )
    @sliderButton.on_click
    def on_sliderButton_clicked(b):
        myRadio = widgets.RadioButtons(
            options=['оставить середину', 'оставить края'],
            layout = { 'width' : '150px'},
            value='оставить середину',
            disabled=False
        )        
        def int_range_slider_widget(wTextarea):

            #joined_string_txt=old_value  
            joined_string_txt=wTextarea.value
            lnames=joined_string_txt.split('\n')
            lnames=[x.strip() for x in lnames]
            lenMax=max(len(w) for w in lnames) #Длина самого длинного файла
            l0=0
            l1=lenMax
            def on_int_range_slider_change(change):
                if myRadio.value=='оставить края':
                    search_input = change['new']
                    vect0 = int(search_input[0])
                    vect1 = int(search_input[1])
                    newList=[]
                    newList.append(old_value)
                    txtA1=[x[:vect0] for x in lnames]
                    txtA2=[x[vect1:] for x in lnames] 
                    joinedA1 = "\n".join(txtA1)
                    joinedA2 = "\n".join(txtA2)
                    txtA=[a+b for i,a in enumerate(txtA1) for j,b in enumerate(txtA2) if i==j]
                    joinedA = "\n".join(txtA)
                    wTextarea.value=joinedA
                if myRadio.value=='оставить середину':
                    search_input = change['new']
                    vect0 = int(search_input[0])
                    vect1 = int(search_input[1])
                    newList=[]
                    newList.append(old_value)
                    txtA3=[x[vect0:vect1] for x in lnames]
                    joinedA3 = "\n".join(txtA3)
                    #txtA=[c for c in enumerate(txtA3)]
                    joinedA="\n".join(txtA3)
                    wTextarea.value=joinedA
            outputSlider = widgets.IntRangeSlider(
                min = l0, 
                step = 1, 
                max = l1, 
                value =[l0,l1])            
            outputSlider.observe(on_int_range_slider_change, 'value')
            return outputSlider    
        def myRadio_on_change(change):
            old_value = allFilesName.value
            txt_value=old_value
            lnames=txt_value.split('\n')
            lnames=[x.strip() for x in lnames]
            lenMax=max(len(w) for w in lnames) #Длина самого длинного файла

            if change['new']=='оставить края':
                l0=lenMax
                l1=lenMax
            else:
                l0=0
                l1=lenMax

            intRangeSlider.min=0
            intRangeSlider.max=l1
            intRangeSlider.value =[l0,l1]
        myRadio.observe(myRadio_on_change, 'value')
        intRangeSlider = int_range_slider_widget(allFilesName) 
        smallHBox = widgets.HBox([intRangeSlider,myRadio])
        sliderButton.disabled=True
        if len(cont1.children)==1:
            cont1.children = (*cont1.children, smallHBox)
        else:
            remove = cont1.children[-1]
            cont1.children = cont1.children[:-1]
            remove.close()
            cont1.children = (*cont1.children, smallHBox)
            
            
#####################################################################
    cont1=widgets.HBox([sliderButton])
    display(cont1)
    display(redactButton)
#######################################################

    ############################################################## def15
    #Кнопка для ексела (загрузить с файла)
    fromExcelButton = widgets.Button(
        description="загрузить имена файлов из ексел",
        layout = { 'width' : 'max-content' },
        disabled = False,
    )
    @fromExcelButton.on_click
    def on_fromExcelButton_clicked(b):
        fileWithNamesExcel='namesToRenameAllFiles.xlsx' #файл в корневой директории. Новые названия файлов
        namesFromExcelToRenameAllFiles(directory, fileWithNamesExcel)    
        
#######################################################
    display(fromExcelButton)
    display(checkButton)
    display(saveButton)
#######################################################

    ############################################################## def16
    #Кнопка для ексела (сохранение)
    toExcelButton = widgets.Button(
        description="сохранить имена файлов в ексел",
        layout = { 'width' : 'max-content' },
        disabled = True,
    )
    @toExcelButton.on_click
    def on_toExcelButton_clicked(b):
        namesFilesToExcel(directory)
        
#############################################################    
    display(toExcelButton)
#######################################################        

###########################################################################################

def natural_sort_with_count_files(directory, count):
    import re
    import os
    def natural_sort(l): 
        convert = lambda text: int(text) if text.isdigit() else text.lower() 
        alphanum_key = lambda key: [ convert(c) for c in 
        re.split('([0-9]+)', key) ] 
        return sorted(l, key = alphanum_key)
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return 
    if count<1:
        return 
    files=natural_sort(files)
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        text=f_name
        pattern=r"^"
        repl=str(ids+1).zfill(count)
        m=re.sub(pattern, repl, text)
        new_name=m.strip()+f_ext
        os.rename(directory+'/'+f, directory+'/'+new_name)

def namesFromExcelToRenameAllFiles(directory, excelFile):
    import os
    import pandas as pd
    import xlsxwriter
    import xlrd
    import openpyxl
    
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return
    
    dfIn = pd.read_excel(excelFile, engine='openpyxl')
    dfName=dfIn['Names Files']
    dfNameList = dfName.values.tolist()
    print (len(dfNameList))
    if len(dfNameList) !=  len(files):
            raise TypeError("Количество имен и файлов НЕ равны")
            return    
    

    for ids,f in enumerate(files):
            f_name, f_ext=os.path.splitext(f)
            f_name_new=str(dfNameList[ids])+f_ext
            os.rename(directory+'/'+f, directory+'/'+f_name_new)

def splitPdfAllPages(pdfFile):
    from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
    import os
    f_name, f_ext=os.path.splitext(pdfFile)

    with open(pdfFile, 'rb') as infile:
        reader = PdfFileReader(infile)
        page = reader.numPages
        for i in range(page):
            writerStart = PdfFileWriter()
            writerStart.addPage(reader.getPage(i))
            fpage=f_name+str(i+1)+f_ext
            print(fpage)
            with open(fpage, 'wb') as outfileStart:
                writerStart.write(outfileStart)
                print(f"Страницу файла сохранил как {fpage}")

def splitPdfTwoPart(pdfFile,numBreak):
    import os
    from PyPDF2 import PdfFileReader, PdfFileWriter, PdfFileMerger
    if ((not isinstance(numBreak, int)) or (numBreak<1)):
            raise TypeError("numBreak должно быть целое число не меньше 1")
            return

    with open(pdfFile, 'rb') as infile:
        reader = PdfFileReader(infile)
        page = reader.numPages
        f_name, f_ext=os.path.splitext(pdfFile)

        if (numBreak>page-1):
            raise TypeError("numBreak должно быть целое число не больше страниц файла")
            return
        writerStart = PdfFileWriter()
        writerEnd = PdfFileWriter()
        for i in range(numBreak):
            writerStart.addPage(reader.getPage(i))
        for i in range(numBreak,page):
            writerEnd.addPage(reader.getPage(i))

        print(f"Всего страниц в файле {pdfFile} - {page}")
        print(f"Разбиваю файл после страницы {numBreak}")

        with open('pagesStart.pdf', 'wb') as outfileStart:
            writerStart.write(outfileStart)
        with open('pagesEnd.pdf', 'wb') as outfileEnd:
            writerEnd.write(outfileEnd)
        print(f"Начало файла сохранил как pagesStart.pdf")
        print(f"Конец файла сохранил как pagesEnd.pdf")

def insertStrForFilesName(directory,string,afterNumber):
    import os
    import re
    if ((not isinstance(afterNumber, int)) or (afterNumber<-1)):
            raise TypeError("afterNumber должно быть целое число не меньше -1")
            return

    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return 
     
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        text=f_name
        pattern=""
        repl=str(string)
        if afterNumber==0:
            pattern=r"^"
        elif afterNumber==-1 or afterNumber>len(f):
            pattern=r"$"
        else:
            pat= '.'*int(afterNumber) 
            pattern=re.match(pat,f_name).group(0)
            repl=pattern+str(string)
        m=re.sub(pattern, repl, text)
        print(f"Old Name {text}")
        print(f"New Name {m.strip()}")
        new_name=m.strip()+f_ext
        os.rename(directory+'/'+f, directory+'/'+new_name)

def indexsForFilesWithRE(directory,count,symbols):
    import os
    import re
    if ((not isinstance(count, int)) or (count<1)):
            raise TypeError("count должно быть целое число не меньше 1")
            return
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return 
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        text=f_name
        pattern=r"^"
        repl=str(ids+1).zfill(count)+symbols
        m=re.sub(pattern, repl, text)
        new_name=m.strip()+f_ext
        os.rename(directory+'/'+f, directory+'/'+new_name)

def subForFilesWithRE(directory,pattern,repl):
    import os
    import re
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")
        return 
    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        text=f_name
        m=re.sub(pattern, repl, text)
        new_name=m.strip()+f_ext
        os.rename(directory+'/'+f, directory+'/'+new_name)

def unionPdfsFiles(directory):
    import os
    from PyPDF2 import PdfFileReader, PdfFileMerger
    pdfs = [f for f in os.listdir(directory) if f.endswith(".pdf")]
    if not pdfs:
        print("В папке нет pdf файлов")
        return    
    merger = PdfFileMerger()
    for pdf in pdfs:
        with open(directory+'/'+pdf, 'rb') as fd:
            merger.append(PdfFileReader(fd))
   
    with open('unionPdfsFiles.pdf','wb') as fout:
        merger.write(fout)
    print("Сохранил файл unionPdfsFiles.pdf в основной директории")

def namesFilesToExcel(directory):
  from PyPDF2 import PdfFileReader, PdfFileWriter
  import os
  import pandas as pd
  import xlsxwriter
  files = os.listdir(directory)
  lNames = [] 
  lExt = []
  for idx, f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    lNames.append(f_name)
    lExt.append(f)
  dfToExcel = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel['Full names'] = lExt
  writer = pd.ExcelWriter("namesFiles" + ".xlsx", engine = 'xlsxwriter')
  dfToExcel.to_excel(writer, startrow = 0, startcol = 0, index = False)
  writer.save()
  print("Сохранил файл namesFiles.xlsx в Основной директории")

def docxToPdf(directory):
  from docx2pdf import convert
  import os
  files = [f for f in os.listdir(directory) if f.endswith(".docx")]
  if not files:
    print("В папке нет docx файлов")
    return    
  for ids,f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    name = '{} {} {}'.format(ids+1,f_name, f_ext)
    print(name)
  convert(directory+'/')
  print("pdf файлы сохранены в папке c docx")

def check_tables(directory, numTable):
  numTable = numTable-1
  from docx import Document
  import os
  import copy
  files = os.listdir(directory)
  listSortedNatural = files
  document_out = Document() # новый документ с таблицами 0 всех файлов
  text1 = 'таблицы {} из всех docx файлов'.format(numTable)
  document_out.add_paragraph(text1)
  for ids, rpd in enumerate(listSortedNatural): 
    f_name, f_ext = os.path.splitext(rpd)
    if f_ext != ".docx":
      print("Не docx file: {}".format(rpd))
      continue
    text2 = 'Файл '+str(ids+1)
    document_out.add_paragraph(text2)
    document_out.add_paragraph(rpd)  #print(rpd)
    doc_document = os.path.join(directory, rpd)
    document = Document(doc_document)
    for k, table in enumerate(document.tables):
      if k == numTable:
        template0 = document.tables[numTable]
        tbl0 = template0._tbl
        new_tbl0 = copy.deepcopy(tbl0)
        paragraph1 = document_out.add_paragraph()
        paragraph1._p.addnext(new_tbl0)
    document_out.add_page_break()
  document_out.save('all_tables.docx')
  print("создал файл all_tables.docx")

def unionLeftRight(fileForUnion):
  import pandas as pd
  df_rup = pd.read_excel(fileForUnion, usecols = "A,B" )
  dfIn = pd.read_excel(fileForUnion)
  dfu = pd.DataFrame
  dfu = dfIn['Left']+" "+dfIn['Right']
  writer = pd.ExcelWriter("unionLeftRightOut" + ".xlsx", engine = 'xlsxwriter')
  dfu.to_excel(writer)
  writer.save()
  print("Создал Файл unionLeftRightOut.xlsx в корневом каталоге Питон")

def fromRupFullNames(fileRUP):
  import pandas as pd
  df_rup = pd.read_excel(fileRUP, sheet_name = 'Компетенции(2)', index_col = None, header = None )
  df_rup = df_rup.dropna(subset = [4])
  df_rup.drop(df_rup.head(2).index,inplace = True) # drop last n rows 
  df_rup.drop(df_rup.tail(12).index,inplace = True) # drop last n rows 
  df_rup.dropna(axis = 'columns',how = 'all', inplace = True)
  df_rup.drop(df_rup.columns[[0,1, 3,5]], axis = 1, inplace = True)
  df_rup.reset_index(drop = True, inplace = True)
  # df_rup.columns = ['code','id','text']
  df_rup = df_rup.rename(columns = {2: "Index", 4: "Name"})
  writer = pd.ExcelWriter("rupFullNames" + ".xlsx", engine = 'xlsxwriter')
  df_rup.to_excel(writer, sheet_name = 'FullName', startrow = 0, startcol = 0, index = False)
  writer.save()
  print("сохранил файл rupFullNames.xlsx на корневой директории py")

def pagesFromPdfToExcel(directoryPdfs):
  from PyPDF2 import PdfFileReader, PdfFileWriter
  import os
  import pandas as pd
  import xlsxwriter
  pdfs = [f for f in os.listdir(directoryPdfs) if f.endswith(".pdf")]
  if not pdfs:
    print("В папке нет pdf файлов")
    return
  lNames = [] 
  lPages = []
  for idx, f in enumerate(pdfs):
    pdf = PdfFileReader(directoryPdfs+"/"+f)
    page = pdf.numPages
    lPages.append(page)
    f_name, f_ext = os.path.splitext(f)
    lNames.append(f_name)
  dfNamesFiles = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel = pd.DataFrame (lNames,columns = ['Names Files'])
  dfToExcel['Кол-во листов'] = lPages
  writer = pd.ExcelWriter("pagesOfPdfFiles" + ".xlsx", engine = 'xlsxwriter')
  dfToExcel.to_excel(writer, startrow = 0, startcol = 0, index = False)
  writer.save()
  print("Сохранил файл pagesOfPdfFiles.xlsx на корневом каталоге py")


def renameAllFiles(directory,inStrins,outString,numStart,numEnd):
    import os
    files = os.listdir(directory)
    if not files:
        print("В папке нет файлов")

    for ids,f in enumerate(files):
        f_name, f_ext=os.path.splitext(f)
        print("Old Name {}".format(f_name))
        f_name=f_name.replace(inStrins,outString).strip()
        f_start=f_name
        if numStart>0:
            f_start=f_name[:numStart]
        f_end=""
        if numEnd>0:
            f_end=f_name[len(f_name)-numEnd:]
        f_name=f_start+f_end
        print("New Name {}".format(f_name))
        new_name='{}{}'.format(f_name, f_ext)
        os.rename(directory+'/'+f, directory+'/'+new_name)


def indexForFiles(directory,constStr,startNum):
  import os
  files = os.listdir(directory)
  for ids,f in enumerate(files):
    f_name, f_ext = os.path.splitext(f)
    print("Old Name {}".format(f_name))
    new_name = '{}{}'.format(f_name, f_ext)
    new_ids = str(ids+startNum)
    new_name = '{}{} {}{}'.format(constStr,new_ids,f_name, f_ext)
    print("New Name {}".format(new_name))
    os.rename(directory+'/'+f, directory+'/'+new_name)
    
def unionLeftRight(fileForUnion):
  import pandas as pd
  import xlsxwriter
  import xlrd
  import openpyxl
  df_rup = pd.read_excel(fileForUnion, usecols = "A,B", engine='openpyxl' )
  dfIn = pd.read_excel(fileForUnion, engine='openpyxl')
  dfu = pd.DataFrame
  dfu = dfIn['Left']+" "+dfIn['Right']
  writer = pd.ExcelWriter("unionLeftRightOut" + ".xlsx", engine = 'xlsxwriter')
  dfu.to_excel(writer)
  writer.save()
  print("Создал Файл unionLeftRightOut.xlsx в корневом каталоге")
